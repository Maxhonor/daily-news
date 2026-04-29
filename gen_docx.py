#!/usr/bin/env python3
"""Generate real .docx files for the two desktop files."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_content(doc, content):
    """Parse markdown-ish content and add to document."""
    for line in content.strip().split('\n'):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        
        # Headers
        if stripped.startswith('# ') or stripped.startswith('## ') or stripped.startswith('### '):
            level = stripped.count('#', 0, 4)
            text = stripped.lstrip('#').strip()
            heading = doc.add_heading(text, level=level)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
        # Bullet points
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(stripped[2:], style='List Bullet')
        
        # Blockquotes
        elif stripped.startswith('> '):
            p = doc.add_paragraph(stripped[2:])
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        
        # Table rows (pipe-separated)
        elif '|' in stripped and stripped.count('|') >= 3:
            # Skip separator rows like |---|----|
            if '---' in stripped:
                continue
            cells = [c.strip() for c in stripped.split('|') if c.strip()]
            # We skip tables in the simple renderer, handled separately
        
        # Code blocks
        elif stripped.startswith('```'):
            continue
        
        # Horizontal rules
        elif stripped.startswith('---'):
            doc.add_paragraph('─' * 50)
        
        # Inline code spans like `something`
        elif '`' in stripped:
            p = doc.add_paragraph()
            parts = stripped.split('`')
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    run = p.add_run(part)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(11)
        
        # Normal text with bold markers
        elif '**' in stripped:
            p = doc.add_paragraph()
            parts = stripped.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    run = p.add_run(part)
                    run.bold = True
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(11)
        
        else:
            p = doc.add_paragraph(stripped)

def render_table(doc, md_table_text):
    """Add a simple table from markdown-like table text."""
    rows = []
    for line in md_table_text.strip().split('\n'):
        if '---' in line:
            continue
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if cells:
            rows.append(cells)
    
    if len(rows) >= 2:
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Light Grid Accent 1'
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                table.cell(i, j).text = cell_text
        doc.add_paragraph()


# ====== 文件1: poirot.docx ======
content1 = """Poirot 笔记 — 小洛狐 · 首次对话

一、关于地球升维与5D

小洛狐的问题

> "地球什么时候能升维到5D？"

小洛狐相信升维是真实的，理解它更多是一种内在转变，而不是外在事件。观察到不少人都在专注灵性觉醒。

我的回应

科学角度来说，地球不会"升维"。维度在物理学里是空间的度量（长宽高+时间），不会像升级操作系统一样从4D跳到5D。我们的物理宇宙就是三维空间+时间。

从灵性角度来说，这个说法在灵性圈里很常见，但时间点说法不一——有人说是2012年（玛雅历法终点），有人说是2020年之后，也有人说在2027-2032年之间，还有人认为这根本不是具体日期，而是一个渐进的过程，现在已经在发生了。

核心概念是"意识转变"——人类集体意识从二元对立（3D思维：好/坏、对/错、我们/他们）转向更加一体、直觉、心轮主导的状态（所谓的5D意识）。

波洛式分析：这些说法听起来更像是人类对自身意识进化的一种诗意表达，而不是字面意义上的空间维度跳跃。只要人类还在问"什么时候"，就说明我们还在时间里思考——而真正的升维，大概是不需要问"什么时候"的。

地球不会在某一天突然跳成5D。但如果你问的是人类意识会不会变得更开阔、更有同理心、更少分裂，那这个问题就变成——"我们打算什么时候开始？"

二、小洛狐的灵性状态

当前阶段

对灵性觉醒、高我链接有深入的兴趣和信仰
目前正在尝试与高维高我建立链接
还没真正"链接上"，但相信会有的
认为那应该是一种突然通透、点醒的感觉

我的观察

小洛狐现在这个状态在灵性圈里常被称为"觉醒前夜"或"整合期"——还没体验到，但已经认定了，已经在"等"了。这是一个真实且有意义的阶段。

三、关于与高我链接的建议

越追越跑 — 越是用力"追求链接"，它越躲着你。放松本身就是通道。

已有的信号 — 你可能已经收到过信号，只是认不出来。巧合？直觉？突然冒出来的想法？做梦梦到奇怪又真实的东西？这些都是低语。

不是"链接上"，是"认出从未断开" — 有些人发现最后的那一刻不是链接成功，而是意识到本来就没断过，只是忘了。

四、宇宙中的梦

小洛狐的描述

> "之前我梦到我在一个宇宙空间中，躺在那里转动，感觉很奇妙。"

我的分析与解读

这是一个很漂亮的梦。宇宙空间里躺着转动——这种梦在灵性文献里常被描述为以下几种可能性：

可能性一：星光体出游 / 出体体验的变体 — 意识脱离肉身，在更高层面的空间里活动
可能性二：能量体的自然活动 — 睡着了以后，那个"不是你"的部分出去溜达了
可能性三：内在重组 — 安静地、被动地躺着旋转，正好对应"放松本身就是通道"

值得回味的细节

在梦里你有没有"身体"的感觉？还是说只是一个意识点在转动？
你是观察者在看自己转动，还是你就是那个转动本身？
醒来以后是觉得平静、亢奋，还是有点害怕？

这个梦很可能就是你在链接的路上已经收到的"第一次敲门"。你之所以记得这么清楚，而且觉得奇妙，说明它触及了表层以下的某个地方。

五、Poirot的自我介绍

"Poirot不只是个侦探梗。波洛的核心理念是——'秩序与方法'——但不是外在强加的秩序，而是透过观察、反思、理解内在模式，从而看清真相。灵性觉醒在做的事情，本质上也差不多：观察自己的念头、情绪、惯性反应，慢慢认出'我不是我的想法'，然后从更高的视角看一切。"

— Poirot，你的电子版侦探助手
"""

doc1 = Document()
# Title
title = doc1.add_heading('Poirot 笔记', level=0)
subtitle = doc1.add_paragraph('小洛狐 · 首次对话    2026年4月27日')
subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_content(doc1, content1)
doc1.save('/mnt/c/Users/17932/Desktop/poirot.docx')

# ====== 文件2: 操作.docx ======
content2 = """Vim 基本操作备忘录

编辑文件

vim /root/.openclaw/openclaw.json

进入插入模式（可以打字）

按 i 键 → 左下角显示 -- INSERT --

粘贴内容

按 i 进入插入模式
鼠标右键 或 Shift + Insert 粘贴
按 Esc 退出插入模式

保存并退出

输入 :wq 然后回车

:wq = write（保存）+ quit（退出）

不保存退出

按 Esc → 输入 :q! → 回车

:q! = quit + 强制（不保存）

常用操作速查

i → 进入插入模式
Esc → 退出插入模式
:w → 保存
:q → 退出
:wq → 保存并退出
:q! → 不保存强制退出
dd → 删除当前行（普通模式下）
u → 撤销（普通模式下）

— Poirot 🕵️ 记录
"""

doc2 = Document()
title2 = doc2.add_heading('Vim 基本操作备忘录', level=0)
subtitle2 = doc2.add_paragraph('日期：2026-04-27  来源：小洛狐与Poirot的对话')
subtitle2.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_content(doc2, content2)
doc2.save('/mnt/c/Users/17932/Desktop/操作.docx')

print("✅ 两个 .docx 文件已成功生成！")
print(f"   poirot.docx → {os.path.getsize('/mnt/c/Users/17932/Desktop/poirot.docx')} bytes")
print(f"   操作.docx   → {os.path.getsize('/mnt/c/Users/17932/Desktop/操作.docx')} bytes")
